from gpt_index import Document, GPTListIndex
import openai
import os
import PyPDF2
import docx
import pytesseract
from PIL import Image
import os

os.environ["OPENAI_API_KEY"] = "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"
openaikey = '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'
import requests

import subprocess

def pdftotext(file_name):
    """
    Function to extract text from .pdf format files
    """

    text = []
    # Open the PDF file in read-binary mode
    with open(file_name, 'rb') as file:
        # Create a PDF object
        pdf = PyPDF2.PdfReader(file)

        # Get the number of pages in the PDF document
        num_pages = len(pdf.pages)

        # Iterate over every page
        for page in range(num_pages):
            # Extract the text from the page
            result = pdf.pages[page].extract_text()
            text.append(result)

    text = "\n".join(text)

    return text


def docxtotext(file_name):
    """
    Function to read .docx format files
    """
    file_name =  file_name
    # Open the Word document
    document = docx.Document(file_name)

    # Extract the text from the document
    text = '\n'.join([paragraph.text for paragraph in document.paragraphs])

    return text


def readtextfile(file_name):
    """
    Function to read .txt format files
    """
    file_name = file_name

    # Open the Text document
    with open(file_name, 'r') as file:
        text = file.read()

    return text


def imagetotext(file_name):
    """
    Function to extract text from images
    """
    file_name = file_name
    # Open the image using PIL
    image = Image.open(file_name)

    # Extract the text from the image
    text = pytesseract.image_to_string(image)

    return text


def preprocesstext(text):
    """
    Function to preprocess text
    """
    # Split the string into lines
    lines = text.splitlines()
    # Use a list comprehension to filter out empty lines
    lines = [line for line in lines if line.strip()]
    # Join the modified lines back into a single string
    text = '\n'.join(lines)

    return text


def processfiles(files, filename):
    """
    Function to extract text from documents
    """
    textlist = []

    # Iterate over provided files
    for file in files:
        # Get file name
        # file_name = file.filename
        file_name = filename
        # Get extention of file name
        ext = file_name.split(".")[-1].lower()

        # Process document based on extention
        if ext == "pdf":
            text = pdftotext(file_name)
        elif ext == "docx":
            text = docxtotext(file_name)
        elif ext == "txt":
            text = readtextfile(file_name)
        elif ext in ["png", "jpg", "jpeg"]:
            text = imagetotext(file_name)
        else:
            text = ""

        # Preprocess text
        text = preprocesstext(text)

        # Append the text to final result
        textlist.append(text)

    return textlist


def createdocuments(textlist):
    """
    Function to create documents as needed for indexing.
    """
    documents = []
    # Create Document for indexing
    for text in textlist:
        documents.append(Document(text))

    return documents


def fileformatvaliditycheck(files, filename):
    """
    Function to check validity of file formats
    """

    for file1 in files:
        # file_name = file1.filename
        file_name = filename
        # Get extention of file name
        ext = file_name.split(".")[-1].lower()

        if ext not in ["pdf", "txt", "docx", "png", "jpg", "jpeg"]:
            return False
    return True


def openaiapikeyvaliditycheck(openaikey):
    """
    Function to check validity of openai key
    """
    # Set the API key
    openai.api_key = openaikey
    # Test the API key by making a request to the OpenAI API
    try:
        response = openai.Model.list()
        # print(response)
        return "Valid OpenAI API key"
    except openai.OpenAIError:
        apikeylink = "https://beta.openai.com/account/api-keys"
        return f"Incorrect OpenAI API key provided: {openaikey}. You can find your OpenAI API key here - {apikeylink}"


def createindex(files,filename, openaikey):
    """
    Function to create index
    """

    # Basic Checks
    if not files:
        return "Upload file before proceeding further."

    fileformatvalidity = fileformatvaliditycheck(files, filename)

    if not fileformatvalidity:
        return "Please upload documents in pdf/txt/docx/png/jpg/jpeg format only."

    if not openaikey:
        return "Please enter your openai key."

    openaiapikeyvality = openaiapikeyvaliditycheck(openaikey)

    if openaiapikeyvality != "Valid OpenAI API key":
        return openaiapikeyvality

    # Store openai key in environment
    os.environ['OPENAI_API_KEY'] = openaikey

    # Process the Documents
    doctextlist = processfiles(files, filename)
    documents = createdocuments(doctextlist)

    # Create index
    index = GPTListIndex(documents, chunk_size_limit=3500)
    # Save index
    # index.save_to_disk('index.json')
    # print(index)
    return index

    return "Uploading documents successfully. OpenAI API Key provided is Valid."


def docques(query, openaikey):
    """
    Function to for quering on the index created
    """
    # Store openai key in environment
    os.environ['OPENAI_API_KEY'] = openaikey

    # Load index
    index = GPTListIndex.load_from_disk('index.json')

    
    # Query based on index
    response = index.query(query, response_mode="tree_summarize")

    return response


def cleartext(query, output):
    """
    Function to clear text
    """
    return ["", ""]

from bs4 import BeautifulSoup
def processurl(url):

    # specify the URL of the webpage you want to extract
    # url = 'https://www.example.com/'

    # send a GET request to the URL
    response = requests.get(url)

    # extract the HTML content from the response object
    html_content = response.content

    # create a Beautiful Soup object and parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # extract all the text from the HTML content
    text = soup.get_text()

    text = preprocesstext(text)
    # print the extracted text
    # print(text)

    textlist = [text]

    return textlist


def createindexfromurl(url, openaikey):
    """
    Function to create index from url
    """
    # Basic Checks
    if not url:
        return "Please enter url before proceeding further."

    if not openaikey:
        return "Please enter your openai key."

    openaiapikeyvality = openaiapikeyvaliditycheck(openaikey)

    if openaiapikeyvality != "Valid OpenAI API key":
        return openaiapikeyvality

    # Store openai key in environment
    os.environ['OPENAI_API_KEY'] = openaikey

    # Process the Documents
    doctextlist = processurl(url)
    documents = createdocuments(doctextlist)

    # Create index
    index = GPTListIndex(documents, chunk_size_limit=3500)
    # Save index
    # index.save_to_disk('index.json')
    # print(index)
    return index

    return "Uploading documents successfully. OpenAI API Key provided is Valid."




def pdf_to_text(file, file_name, query):
    # Save the file
    # file_path =  file_name
    # with open(file_path, 'wb') as f_out:
    #     f_out.write(file.read())

    print("====Generating index====")
    index = createindex([file], file_name, openaikey)
    index.save_to_disk('index.json')

    print("====querying====")
    response = docques(query, openaikey)

    # print(response)
    return response.__str__()


def pdf_to_commands(file_path, query):
    file_name = file_path.split('/')[-1]
    with open(file_path, 'rb') as f:
        # print(f.read())
        response = pdf_to_text(f, file_name, query)
    print("==== Commands Extracted ====")
    print(response)
    return response


def download_repo_and_move_to_dir(repo_url):
    process = subprocess.Popen(f'git clone {repo_url}', shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    stdout, stderr = process.communicate()

    if process.returncode == 0:
        print(stdout.decode('utf-8'))
    else:
        print(f'Error: {stderr.decode("utf-8")}')
        
    repo_name = repo_url.split('/')[-1].split('.')[0]
    os.chdir(repo_name)
    return repo_name

def get_commands(commands):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "i will give you setup commands and need to only extract the commands in a list format and do modify and add and remove some commands like git clone and cd move/to/dir commands because i have already downloaded the repo and inside the repo there should be no errors in the commands and nothing else except a list of commands"},
            {"role": "user", "content": """
                The setup commands are:
                    1. `git clone https://github.com/wtran29/Blackjack-Tkinter.git`
                    2. `cd path/to/Blackjack-Tkinter`
                    3. `python3 blackjack.py`
                    """},
            {"role": "assistant", "content": "['python3 blackjack.py']"},
            {
                "role": "user", "content": """
                The setup commands are: 
                    1. `pip install -r requirement.txt` 
                    2. `python app.py`
                    """
            },
            {"role": "assistant", "content": "['pip install -r requirement.txt', 'python app.py']"},
            {
                "role": "user", "content": """
                1. Download the pre-trained model from [here](https://drive.google.com/file/d/0B5MzpY9kBtDVZ2RpVDYwWmxoSUk/edit) and move the model file to the 'model/' folder, the path of the model should be as follows: ```mv <downloaded_model_file> model/20170512-110547/20170512-110547.pb```
                2. Install Python 3.6
                3. Install CUDA Toolkit 9.0
                4. Install cuDNN 7.1.4
                5. Install the required libraries for running on CPU: ```pip3 install -r requirements_cpu.txt```
                6. Install the required libraries for running using a CUDA GPU: ```pip3 install -r requirements_gpu.txt```
                7. Run the server by using the ```python server.py``` command.
                """
            },
            {
                "role": "user", "content": commands
            }
        ]
    )

    return response


def run_commands(commands):
    
    for command in commands:
        if command.startswith('cd '):
            # Change the working directory for the Python script
            os.chdir(command[3:])
        else:
            process = subprocess.Popen(command, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = process.communicate()

            if process.returncode == 0:
                print(stdout.decode('utf-8'))
            else:
                print(f'Error: {stderr.decode("utf-8")}')
        

import argparse

def convert_md_to_txt(file_path):
    with open(file_path, "rb") as f:
        content = f.read()
    # print(content)
    txt_file_path = file_path.replace(".md", ".txt")
    with open(txt_file_path, "wb") as f:
        f.write(content)
    return txt_file_path

def main(git_url):
    download_repo_and_move_to_dir(git_url)
    md_file_path = "README.md"
    txt_file_path = convert_md_to_txt(md_file_path)
    response = pdf_to_commands(txt_file_path, "Extract all the necessary commands that needs to be run after cloning the repo including any download links of files if there is any in order to run the project")
    commands = get_commands(response)
    command_list = commands.choices[0].message.content

    clean_string = command_list.strip("[]").replace("'", "")
    # Split the cleaned string using the delimiter ', '
    result_list = clean_string.split(', ')
    print(result_list)



if __name__ == "__main__":
    # Parse command line arguments
    parser = argparse.ArgumentParser(description="A command line tool for downloading a repo and processing a README.")
    parser.add_argument("git_url", help="The URL of the Git repository to download.")
    args = parser.parse_args()

    # Call the main function with the provided git_url
    main(args.git_url)