import datetime
import pandas as pd
import numpy as np
from hbshare.fe.XZ import db_engine
from hbshare.fe.XZ import functionality
from hbshare.fe.mutual_analysis import pool_analysis as pla




util=functionality.Untils()
hbdb=db_engine.HBDB()
localdb=db_engine.PrvFunDB().engine

jjdm_list = util.get_mutual_stock_funds(datetime.datetime.today().strftime('%Y%m%d'))
jjdm_con = util.list_sql_condition(jjdm_list)
sql = "select jjjc,jjdm from st_fund.t_st_gm_jjxx where jjdm in ({0})".format(jjdm_con)
jjdm_name_map = hbdb.db2df(sql, db='funduser')

class Equilibrium:

    @staticmethod
    def ind_equ():

        latest_asofdate=\
            pd.read_sql("select max(asofdate) as asofdate from hbs_industry_property_new ",con=localdb)['asofdate'][0]

        sql="SELECT jjdm,cen_ind_1,asofdate from hbs_industry_property_new where asofdate='{0}' order by cen_ind_1 ".format(latest_asofdate)
        ind_equ=pd.read_sql(sql,con=localdb)

        return ind_equ

    @staticmethod
    def style_equ():

        latest_asofdate = \
        pd.read_sql("select max(asofdate) as asofdate from hbs_style_property ", con=localdb)['asofdate'][0]

        sql="""SELECT jjdm,cen_lv_rank,asofdate 
        from hbs_style_property where asofdate='{0}' order by cen_lv_rank 
        """.format(latest_asofdate)
        style_equ = pd.read_sql(sql, con=localdb)

        return  style_equ

    @staticmethod
    def size_equ():

        latest_asofdate = \
        pd.read_sql("select max(asofdate) as asofdate from hbs_size_property ", con=localdb)['asofdate'][0]

        sql = """SELECT jjdm,cen_lv_rank,asofdate 
        from hbs_size_property where asofdate='{0}' order by cen_lv_rank 
        """.format(latest_asofdate)

        size_equ = pd.read_sql(sql, con=localdb)

        return  size_equ

    @staticmethod
    def pepbroe_equ():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_holding_property ", con=localdb)[
                'asofdate'][0]

        sql="""
        select jjdm,PE_rank,PB_rank,ROE_rank,股息率_rank,asofdate 
        from hbs_holding_property where asofdate='{0}'
        """.format(latest_asofdate)

        pepbroe_equ=pd.read_sql(sql,con=localdb)
        pepbroe_equ[
            ['PE_rank','PB_rank','ROE_rank','股息率_rank']
        ]=abs(pepbroe_equ[['PE_rank','PB_rank','ROE_rank','股息率_rank']]-0.5)

        pepbroe_equ=pepbroe_equ.sort_values('PE_rank')

        return pepbroe_equ

    @staticmethod
    def nav_equ():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from nav_ret_bias ", con=localdb)[
                'asofdate'][0]

        sql="""
        select jjdm,mean_rank_monthly,std_rank_monthly,mean_rank_yearly,std_rank_yearly,asofdate
        from nav_ret_bias where asofdate='{0}' and mean_rank_yearly is not null  order by mean_rank_monthly
        """.format(latest_asofdate)

        nav_equ=pd.read_sql(sql,con=localdb)

        return  nav_equ

    def get_equilibrium(self,threshield,show_num=300):

        # method_1
        ind_equ = self.ind_equ()
        style_equ = self.style_equ()
        size_equ = self.size_equ()
        pepbroe_equ = self.pepbroe_equ()
        nav_equ = self.nav_equ()

        #rename columns :
        ind_equ=pd.merge(ind_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').\
            rename(columns={'cen_ind_1':'行业集中度'})

        style_equ=pd.merge(style_equ,
                         jjdm_name_map,
                         how='left', on='jjdm')\
            .rename(columns={'cen_lv_rank': '风格集中度'})

        size_equ=pd.merge(size_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').\
            rename(columns={'cen_lv_rank': '规模集中度'})

        pepbroe_equ=pd.merge(pepbroe_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').rename(columns={'PE_rank': 'pe偏离度',
                                    'PB_rank': 'pb偏离度',
                                    'ROE_rank': 'roe偏离度',
                                    '股息率_rank': '股息率偏离度'})

        nav_equ=pd.merge(nav_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').rename(columns={'mean_rank_monthly': '净值偏离度（月）',
                                'std_rank_monthly': '净值偏离度方差（月）',
                                'mean_rank_yearly': '净值偏离度（年）',
                                'std_rank_yearly': '净值偏离度方差（年）',
                                })

        # method_2
        joint_rank = pd.merge(ind_equ, style_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, size_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, pepbroe_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, nav_equ, how='inner', on=['jjdm','jjjc'])

        col_list = ['行业集中度', '风格集中度',
                    '规模集中度', 'pe偏离度',
                    '净值偏离度（月）']

        joint_rank['平均集中度'] = joint_rank[col_list].mean(axis=1)
        joint_rank = joint_rank.sort_values('平均集中度')

        # method_3
        joint_result = joint_rank[(joint_rank[col_list] <= threshield).prod(axis=1) == 1][['jjdm','jjjc','asofdate'] + col_list]

        joint_rank=joint_rank[['jjdm','jjjc','平均集中度','asofdate']+col_list]


        #format change
        ind_equ['行业集中度'] = ind_equ['行业集中度'].map("{:.2%}".format)
        style_equ['风格集中度'] = style_equ['风格集中度'].map("{:.2%}".format)
        size_equ['规模集中度'] = size_equ['规模集中度'].map("{:.2%}".format)

        for col in ['pe偏离度','pb偏离度','roe偏离度','股息率偏离度']:
            pepbroe_equ[col] = pepbroe_equ[col].map("{:.2%}".format)
        for col in ['净值偏离度（月）', '净值偏离度方差（月）', '净值偏离度（年）', '净值偏离度方差（年）']:
            nav_equ[col] = nav_equ[col].map("{:.2%}".format)
        for col in col_list:
            joint_rank[col]=joint_rank[col].map("{:.2%}".format)
            joint_result[col]=joint_result[col].map("{:.2%}".format)
        joint_rank['平均集中度']=joint_rank['平均集中度'].map("{:.2%}".format)

        return ind_equ[0:show_num],style_equ[0:show_num],\
               size_equ[0:show_num],pepbroe_equ[0:show_num],nav_equ[0:show_num],\
               joint_rank[0:show_num],joint_result[0:show_num]

class Leftside:

    @staticmethod
    def stock_left():

        latest_asofdate=\
            pd.read_sql("select max(asofdate) as asofdate from hbs_stock_trading_property ",
                        con=localdb)['asofdate'][0]

        sql="""
        SELECT jjdm,`左侧概率（出持仓前,半年线）_rank`,`左侧概率（出持仓前,年线）_rank`,asofdate 
        from hbs_stock_trading_property where asofdate='{0}'  """\
            .format(latest_asofdate)
        stock_left=pd.read_sql(sql,con=localdb)

        stock_left['stock_left_rank']=stock_left[['左侧概率（出持仓前,半年线）_rank','左侧概率（出持仓前,年线）_rank']]\
            .max(axis=1)

        stock_left=stock_left.sort_values('stock_left_rank',ascending=False)

        return stock_left

    @staticmethod
    def ind_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_industry_shift_property_new "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_industry_shift_property_new where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        ind_left=pd.read_sql(sql,con=localdb)

        ind_left=ind_left.groupby('jjdm').max('Total_rank')
        ind_left.reset_index(inplace=True)
        ind_left['asofdate']=latest_asofdate

        ind_left.sort_values('Total_rank',ascending=False,inplace=True)

        ind_left.rename(columns={'Total_rank':'ind_rank'},inplace=True)

        return ind_left

    @staticmethod
    def value_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_shift_property_value "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_shift_property_value where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        value_left=pd.read_sql(sql,con=localdb)

        value_left=value_left.groupby('jjdm').max('Total_rank')
        value_left.reset_index(inplace=True)
        value_left['asofdate']=latest_asofdate

        value_left.sort_values('Total_rank',ascending=False,inplace=True)
        value_left.rename(columns={'Total_rank': 'value_rank'}, inplace=True)

        return value_left

    @staticmethod
    def size_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_shift_property_size "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_shift_property_size where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        size_left=pd.read_sql(sql,con=localdb)

        size_left=size_left.groupby('jjdm').max('Total_rank')
        size_left.reset_index(inplace=True)
        size_left['asofdate']=latest_asofdate

        size_left.sort_values('Total_rank',ascending=False,inplace=True)
        size_left.rename(columns={'Total_rank': 'size_rank'}, inplace=True)


        return size_left

    def get_left(self,threshield,show_num=300):

        # method_1
        stock_left = self.stock_left()
        ind_left = self.ind_left()
        value_left = self.value_left()
        size_left = self.size_left()

        #rename columns
        stock_left=stock_left[['jjdm','stock_left_rank','asofdate']]

        stock_left=pd.merge(stock_left,jjdm_name_map,how='left',on='jjdm')\
            .rename(columns={'stock_left_rank':'个股交易左侧率'})
        ind_left = pd.merge(ind_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'ind_rank':'行业切换左侧率'})
        value_left = pd.merge(value_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'value_rank':'风格切换左侧率'})
        size_left = pd.merge(size_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'size_rank':'规模切换左侧率'})


        # method_2
        joint_rank = pd.merge(stock_left, ind_left, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, value_left, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, size_left, how='inner', on=['jjdm','jjjc'])
        col_list = ['个股交易左侧率', '行业切换左侧率', '风格切换左侧率', '规模切换左侧率']
        joint_rank['平均左侧率'] = joint_rank[col_list].mean(axis=1)

        # method_3
        joint_restult = joint_rank[(joint_rank[col_list] >= (1 - threshield)).prod(axis=1) == 1][['jjdm','jjjc'] + col_list]

        joint_rank = joint_rank.sort_values('平均左侧率', ascending=False)[['jjdm','jjjc', '平均左侧率', 'asofdate_x']+col_list]


        #re format
        stock_left['个股交易左侧率']=stock_left['个股交易左侧率'].map("{:.2%}".format)
        ind_left['行业切换左侧率'] = ind_left['行业切换左侧率'].map("{:.2%}".format)
        value_left['风格切换左侧率'] = value_left['风格切换左侧率'].map("{:.2%}".format)
        size_left['规模切换左侧率'] = size_left['规模切换左侧率'].map("{:.2%}".format)
        for col in col_list:
            joint_rank[col] = joint_rank[col].map("{:.2%}".format)
            joint_restult[col] = joint_restult[col].map("{:.2%}".format)

        joint_rank['平均左侧率']=joint_rank['平均左侧率'].map("{:.2%}".format)

        return stock_left[0:show_num],ind_left[0:show_num],value_left[0:show_num],size_left[0:show_num],joint_rank[0:show_num],joint_restult[0:show_num]

class Size:
    @staticmethod
    def size_property(fre):
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_size_property".format(fre), con=localdb)['asofdate'][0]
        sql = "select jjdm, shift_lv, cen_lv, shift_lv_rank, cen_lv_rank,大盘, 中盘, 小盘 ,大盘_rank, 中盘_rank, 小盘_rank, asofdate from hbs_size_property where asofdate='{0}'".format( latest_asofdate)
        size_property = pd.read_sql(sql, con=localdb)
        return size_property

    def get_size(self, fre, show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5):
        size_property = self.size_property(fre)
        size_property.columns = ['jjdm', '换手率', '集中度', '换手率排名', '集中度排名','大盘(绝对值）', '中盘(绝对值）', '小盘(绝对值）' ,'大盘', '中盘', '小盘', 'asofdate']
        size_property = size_property.merge(jjdm_name_map, on=['jjdm'], how='left')

        size = size_property[(size_property['换手率排名'] < shift_ratio_threshold) & (size_property['集中度排名'] > centralization_threshold)]
        big_size = size[(size['大盘'] > size['中盘']) & (size['大盘'] > size['小盘'])].sort_values('大盘', ascending=False)
        medium_size = size[(size['中盘'] > size['大盘']) & (size['中盘'] > size['小盘'])].sort_values('中盘', ascending=False)
        small_size = size[(size['小盘'] > size['大盘']) & (size['小盘'] > size['中盘'])].sort_values('小盘', ascending=False)

        big_size = big_size[['jjdm', 'jjjc', 'asofdate', '大盘']]
        medium_size = medium_size[['jjdm', 'jjjc', 'asofdate', '中盘']]
        small_size = small_size[['jjdm', 'jjjc', 'asofdate', '小盘']]
        big_size['大盘'] = big_size['大盘'].map("{:.2%}".format)
        medium_size['中盘'] = medium_size['中盘'].map("{:.2%}".format)
        small_size['小盘'] = small_size['小盘'].map("{:.2%}".format)
        return big_size[0: show_num], medium_size[0: show_num], small_size[0: show_num]

class Value:
    @staticmethod
    def value_property():
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from jjpic_value_p_hbs ", con=localdb)['asofdate'][0]
        sql = "SELECT jjdm,`集中度排名(持仓)`,`换手率排名(持仓)`,`成长绝对暴露(持仓)`,`价值绝对暴露(持仓)`,asofdate from jjpic_value_p_hbs where  asofdate='{0}'".format(latest_asofdate)
        value_property = pd.read_sql(sql, con=localdb)
        return value_property

    @staticmethod
    def holding_property():
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_holding_property", con=localdb)['asofdate'][0]
        sql = "select jjdm, PE_rank, PB_rank, PE_REL_rank, PB_REL_rank, ROE_rank, 股息率_rank, asofdate from hbs_holding_property where asofdate='{0}'".format(latest_asofdate)
        holding_property = pd.read_sql(sql, con=localdb)
        return holding_property

    def get_value(self, show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5):

        value_property = self.value_property()
        value_property.columns = ['jjdm', '集中度排名','换手率排名','成长', '价值','asofdate']

        #filter the industry and theme fund

        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from jjpic_theme_p ", con=localdb)['asofdate'][0]
        sql = "SELECT * from jjpic_theme_p where  asofdate='{0}'".format(latest_asofdate)
        jj_theme = pd.read_sql(sql, con=localdb)
        jj_theme['max_theme_w']=jj_theme[[ '大金融', '消费', 'TMT', '周期','制造']].max(axis=1)
        jj_theme=jj_theme[jj_theme['max_theme_w']<=0.7][['jjdm']]

        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from jjpic_industry_detail_1 ", con=localdb)['asofdate'][0]
        sql = "SELECT jjdm,max(`占持仓比例(时序均值)`) as max_ind_w from jjpic_industry_detail_1 where asofdate='{0}' GROUP BY jjdm ".format(latest_asofdate)
        jj_max_ind_w = pd.read_sql(sql, con=localdb)
        jj_max_ind_w=jj_max_ind_w[jj_max_ind_w['max_ind_w']<=0.5][['jjdm']]

        value_property=pd.merge(value_property,jj_theme,how='inner',on='jjdm')
        value_property = pd.merge(value_property, jj_max_ind_w, how='inner', on='jjdm')

        value_property = value_property.merge(jjdm_name_map, on=['jjdm'], how='left')

        #value = value_property[(value_property['换手率排名'] < shift_ratio_threshold) & (value_property['集中度排名'] > centralization_threshold)]
        growth = value_property[value_property['成长'] > value_property['价值']].sort_values('成长', ascending=False)
        value = value_property[value_property['价值'] > value_property['成长']].sort_values('价值', ascending=False)
        growth=growth[(growth['换手率排名']<growth['换手率排名'].quantile(shift_ratio_threshold))]
        value = value[(value['换手率排名'] < value['换手率排名'].quantile(shift_ratio_threshold))]


        growth = growth[['jjdm', 'jjjc', 'asofdate', '成长']]
        value = value[['jjdm', 'jjjc', 'asofdate', '价值']]
        growth['成长'] = growth['成长'].map("{:.2%}".format)
        value['价值'] = value['价值'].map("{:.2%}".format)

        holding_property = self.holding_property()
        holding_property.columns = ['jjdm', 'PE排名', 'PB排名', 'PE相对行业均值排名', 'PB相对行业均值排名', 'ROE排名', '股息率排名', 'asofdate']
        value_holding_property = value.merge(holding_property.drop('asofdate', axis=1), on=['jjdm'], how='left')
        absolute_pe_value = value_holding_property.sort_values('PE排名')
        absolute_pb_value = value_holding_property.sort_values('PB排名')
        relative_pe_value = value_holding_property.sort_values('PE相对行业均值排名')
        relative_pb_value = value_holding_property.sort_values('PB相对行业均值排名')
        dividend_value = value_holding_property.sort_values('股息率排名', ascending=False)
        reverse_value = value_holding_property.sort_values('ROE排名')
        high_quality_value = value_holding_property.sort_values('ROE排名', ascending=False)

        absolute_pe_value = absolute_pe_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PE排名']]
        absolute_pb_value = absolute_pb_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PB排名']]
        relative_pe_value = relative_pe_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PE相对行业均值排名']]
        relative_pb_value = relative_pb_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PB相对行业均值排名']]
        dividend_value = dividend_value[['jjdm', 'jjjc', 'asofdate', '价值', '股息率排名']]
        reverse_value = reverse_value[['jjdm', 'jjjc', 'asofdate', '价值', 'ROE排名']]
        high_quality_value = high_quality_value[['jjdm', 'jjjc', 'asofdate', '价值', 'ROE排名']]

        # absolute_pe_value['PE排名'] = absolute_pe_value['PE排名'].map("{:.2%}".format)
        # absolute_pb_value['PB排名'] = absolute_pb_value['PB排名'].map("{:.2%}".format)
        # relative_pe_value['PE相对行业均值排名'] = relative_pe_value['PE相对行业均值排名'].map("{:.2%}".format)
        # relative_pb_value['PB相对行业均值排名'] = relative_pb_value['PB相对行业均值排名'].map("{:.2%}".format)
        # dividend_value['股息率排名'] = dividend_value['股息率排名'].map("{:.2%}".format)
        # reverse_value['ROE排名'] = reverse_value['ROE排名'].map("{:.2%}".format)
        # high_quality_value['ROE排名'] = high_quality_value['ROE排名'].map("{:.2%}".format)
        return growth[0: show_num], value[0: show_num]
            # , absolute_pe_value[0: show_num], absolute_pb_value[0: show_num], \
            #    relative_pe_value[0: show_num], relative_pb_value[0: show_num], dividend_value[0: show_num], \
            #    reverse_value[0: show_num], high_quality_value[0: show_num]

class Industry_rotating:

    def get_industry_rotating(self,shift_count_rank_thresheld=0.5
                              ,shift_ret_thresheld=0.5,shift_winpro_thresheld=0.5,left_trade_ratio=0.5,pool_size=30):

        max_asofdate=pd.read_sql("select max(asofdate) as max_date from jjpic_industry_sp",con=localdb)['max_date'][0]

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='切换次数' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_count_rank_thresheld,max_asofdate)
        jjdm_list1=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='胜率（直到下次切换）' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_winpro_thresheld,max_asofdate)
        jjdm_list2=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='胜率（下季度）' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_winpro_thresheld,max_asofdate)
        jjdm_list22=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        max_asofdate2 = pd.read_sql("select max(asofdate) as max_date from jjpic_industry_p", con=localdb)['max_date'][0]
        sql="SELECT jjdm from jjpic_industry_p where `一级行业类型` in ('轮动','博弈') and asofdate={0}"\
            .format(max_asofdate2)
        jjdm_list3=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='持有平均收益' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_ret_thresheld,max_asofdate)
        jjdm_list4=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='下季平均收益' and Total_rank>={0} and asofdate='{1}'"\
            .format(shift_ret_thresheld,max_asofdate)
        jjdm_list44=pd.read_sql(sql,con=localdb)['jjdm'].tolist()

        sql="SELECT jjdm from jjpic_industry_sp  where `项目名`='左侧比率' and Total_rank>={0} and asofdate='{1}'"\
            .format(left_trade_ratio,max_asofdate)
        jjdm_list5=pd.read_sql(sql,con=localdb)['jjdm'].tolist()


        jjdm_list=set(jjdm_list1).intersection(set(jjdm_list22))
        jjdm_list=jjdm_list.intersection(set(jjdm_list3))
        jjdm_list = jjdm_list.intersection(set(jjdm_list44))
        jjdm_list=list(jjdm_list.intersection(set(jjdm_list5)))

        sql="select jjdm,avg(Total_rank) as new_factor from jjpic_industry_sp where (`项目名`='下季平均收益' or `项目名`='胜率（下季度）') and jjdm in ({0}) and asofdate='{1}' group by jjdm order by new_factor DESC "\
            .format(util.list_sql_condition(jjdm_list),max_asofdate)
        pool_funds=pd.read_sql(sql,con=localdb).iloc[0:pool_size]['jjdm'].tolist()

        # sql = "select jjdm,info_ratio from factor_infor_ratio where jjdm in ({0}) and date>='{1}' and date <='{2}' and zblb='{3}' and info_ratio!=99999 order by info_ratio DESC" \
        #     .format(util.list_sql_condition(jjdm_list)
        #             , max_asofdate[0:6] + '01', max_asofdate[0:6] + '31'
        #             , '2101')
        # pool_funds = pd.read_sql(sql, con=localdb).iloc[0:pool_size]['jjdm'].tolist()

        sql = "select tjyf,avg(hb1y) from st_fund.t_st_gm_yhb where jjdm in ({0}) and hb1y!=99999 and tjyf>'202112' and tjyf<='202208' group by tjyf " \
            .format(util.list_sql_condition(pool_funds))
        navdf = hbdb.db2df(sql, db='funduser').set_index('tjyf')
        navdf=navdf/100+1
        print((navdf.cumprod()['avg(hb1y)'].iloc[-1]-1)*100)


        return  pool_funds

class Positive_equilibrium:

    def get_nav_equilibrium_pool(self,asofdate,poo_size=30):

        jjdm_list_0=util.get_bmk_funds_list(asofdate)

        sql="""SELECT jjdm,mean_rank_yearly,mean_rank_monthly from nav_ret_bias where mean_yearly is not null 
        and mean_rank_yearly<=0.1 and mean_rank_monthly<=0.1 and asofdate='{0}' and jjdm in ({1}) ORDER BY mean_yearly
        """.format(asofdate[0:6],util.list_sql_condition(jjdm_list_0))
        jjdm_list=pd.read_sql(sql,con=localdb)

        jjdm_list=jjdm_list.iloc[0:poo_size]['jjdm'].tolist()

        return jjdm_list


class Report:


    @staticmethod
    def get_silly_summary(doc,equ_list,equ_name_list,fold,
                          paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num):

        from PIL import Image
        from docx.shared import Inches, Pt


        def style_analysis(df,colname):
            style_style=''
            for style in df.index:
                style_style+=str(style)+"占比"+str(round(df.loc[style][colname],2))+"%,"

            return  style_style

        pic_name_list = ['池列表.png',
                         '行业风格分布.png', '1级行业持仓分布_占持仓比例.png',
                         '1级行业持仓分布_占持仓比例排名.png',
                         '2级行业持仓分布_占持仓比例.png',
                         '2级行业持仓分布_占持仓比例排名.png',
                         '3级行业持仓分布_占持仓比例.png',
                         '3级行业持仓分布_占持仓比例排名.png',
                         '风格类型分布.png',
                         '风格专注型分布.png',
                         '成长价值持仓占比.png',
                         '成长价值持仓占比排名.png',
                         '规模风格类型分布.png',
                         '专注型风格分布.png',
                         '个股持仓特征.png',
                         '大中小盘持仓占比.png',
                         '大中小盘持仓占比排名.png',
                         '个股风格类型分布A.png',
                         '个股风格类型分布B.png',
                         '左侧类型分布.png',
                         '新股偏好分布.png',
                         ]

        for i in range(len(equ_list)):
            equ=equ_list[i]
            name=equ_name_list[i]

            doc.add_paragraph("")
            paragraphs_count += 1

            p=doc.add_paragraph("{2}.{3} {0}：对{1}排名前30的基金进行了统计画像，统计结果如下：".format(name, name,title_num,i+1))
            p.style.font.size=Pt(10)
            p.paragraph_format.first_line_indent=p.style.font.size*2
            p.paragraph_format.line_spacing = Pt(20)  # 行距
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count+=1


            industry_style, class_list, style_type_dis, style_incline_dis, \
            style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
            stock_style_a, stock_style_b, stock_left, stock_new, stock_fin = pla.pool_picturing(equ, returndata=True,save_local_file=True)

            for pic in pic_name_list:
                file_path=fold+"\\"+pic
                try:
                    doc.add_picture(file_path, width=Inches(3), height=Inches(2.8))
                    pic_paragraphs_list.append(paragraphs_count)
                    paragraphs_count += 1
                except Exception as e:
                    pic_temp = Image.open(file_path)
                    pic_temp.save(pic_temp)
                    doc.add_picture(file_path, width=Inches(4.5), height=Inches(2.5))

            ind_style=style_analysis(industry_style,'num')
            style_style=style_analysis(style_type_dis,'风格类型')
            style_weight=style_analysis(style_weight_dis,'成长价值权重分布')
            size_style=style_analysis(style_type_dis2,'规模风格类型')
            size_weight=style_analysis(style_weight_dis2,'成长价值权重分布')
            left_property=style_analysis(stock_left,'左侧类型分布')
            new_stock_property=style_analysis(stock_new,'新股偏好分布')
            stock_hld_property=style_analysis(stock_fin,'个股持仓特征')

            summary="总结来看：从行业上看，{0}池中".format(name)+ind_style+\
                    "一级行业中，{0}占比较高".format(util.list_sql_condition(class_list[0].index[0:3].tolist()))+\
                    "从风格上看，{0}池中".format(name)+style_style+"具体的持仓来看"+style_weight+"。"+\
                    "从规模上看，{0}池中".format(name) + size_style + "具体的持仓来看" + size_weight + "。"+\
                    "从左侧属性上看，{0}池中".format(name) + left_property +"。"+\
                    "从新股属性上看，{0}池中".format(name) + new_stock_property + "。" +\
                    "从个股持仓属性上看，{0}池中".format(name) + stock_hld_property

            doc.add_paragraph("")
            paragraphs_count += 1

            p=doc.add_paragraph(summary)
            p.style.font.size = Pt(10)
            p.paragraph_format.first_line_indent=p.style.font.size*2
            p.paragraph_format.line_spacing = Pt(20)  # 行距
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count+=1

        return doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list

    def filter_pool_report(self):

        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        #save png 2 doc

        equclass=Equilibrium()
        ind_equ, style_equ, size_equ,pepbroe_equ, nav_equ,\
        joint_rank,joint_result=equclass.get_equilibrium(threshield=0.3,show_num=30)
        equ_list=[joint_rank,joint_result]
        equ_name_list=['均衡平均','均衡交集']

        leftclass=Leftside()
        stock_left, ind_left, value_left,size_left,\
        joint_rank, joint_restult=leftclass.get_left(threshield=0.4,show_num=30)
        left_list=[joint_rank, joint_restult]
        left_name_list = ['左侧平均', '左侧交集']

        sizeclass = Size()
        size_list=sizeclass.get_size(fre='M', show_num=30,
                                                               shift_ratio_threshold=0.5,
                                                               centralization_threshold=0.5)
        size_name_list = ['大盘', '中盘', '小盘']

        valueclass = Value()
        value_list=valueclass.get_value(fre='M', show_num=30,
                                        shift_ratio_threshold=0.5, centralization_threshold=0.5)
        value_list=value_list[0:2]
        value_name_list = ['成长', '价值']

        doc = Document()
        fold=r"E:\GitFolder\hbshare\fe\mutual_analysis"

        pic_paragraphs_list=[]
        string_paragraphs_list=[]
        paragraphs_count=0

        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc,equ_list,equ_name_list,fold
                               ,paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=2)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, left_list, left_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=3)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, size_list, size_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=4)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, value_list, value_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=5)

        for j in pic_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for j in string_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


        doc.save(r"E:\GitFolder\hbshare\fe\mutual_analysis\筛选报告\筛选池画像报告.docx")


        print('')

class Advance_man:

    @staticmethod
    def advance_man_from3ind(asofdate):

        advance_man_3ind_cluster=pd.read_excel(r"E:\GitFolder\docs\先进制造三级行业名单20220823.xlsx")[['板块','申万三级']]
       #  target_ind=['电池化学品', '锂电池', '锂电专用设备', '燃料电池', '输变电设备', '风电零部件', '风电整机',
       # '光伏电池组件', '光伏辅材', '光伏加工设备', '硅料硅片', '逆变器', '综合电力设备商', '半导体材料',
       # '半导体设备', '分立器件', '集成电路封测', '集成电路制造', '模拟芯片设计', '数字芯片设计', '品牌消费电子',
       # '消费电子零部件及组装', '印制电路板', '航海装备Ⅲ', '航空装备Ⅲ', '航天装备Ⅲ', '军工电子Ⅲ',
       # '工程机械器件', '工程机械整机', '轨交设备Ⅲ', '能源及重型设备', '工控设备', '机器人', '激光设备',
       # '安防设备', '冰洗', '空调', '厨房电器', '彩电', '清洁小家电', '电动乘用车', '综合乘用车',
       # '车身附件及饰件', '底盘与发动机系统', '轮胎轮毂', '其他汽车零部件', '汽车电子电气系统', '通信网络设备及器件',
       # '通信终端及配件', '体外诊断', '医疗耗材', '医疗设备']

        target_ind=advance_man_3ind_cluster['申万三级'].tolist()

        sql="SELECT jjdm,`占持仓比例(时序均值)`,`行业名称` from jjpic_industry_detail_3 where `行业名称` in({0}) and asofdate='{1}' "\
            .format(util.list_sql_condition(target_ind),asofdate)
        adv_man=pd.read_sql(sql,con=localdb)
        adv_man=pd.merge(adv_man,
                         advance_man_3ind_cluster,
                         how='left',left_on='行业名称',right_on='申万三级').drop('申万三级',axis=1)
        ind3=adv_man.groupby('jjdm').sum().rename(columns={'占持仓比例(时序均值)':'先进制造三级行业'})
        for clu in adv_man['板块'].unique():
            ind3=pd.merge(ind3,adv_man[adv_man['板块']==clu].groupby('jjdm').sum()
                          .rename(columns={'占持仓比例(时序均值)':'{}板块'.format(clu)}),how='inner',on='jjdm')

        tempdf=adv_man.pivot_table('占持仓比例(时序均值)','jjdm','行业名称').fillna(0)
        tempdf.columns=[x+'_w' for x in tempdf.columns.tolist()]
        ind3=pd.merge(ind3,tempdf
                      ,how='left',on='jjdm')

        ind3=ind3.sort_values('先进制造三级行业',ascending=False)

        return  ind3
    @staticmethod
    def advance_man_wording():

        key_words=['先进制造','高端制造','智能制造','国产替代','专精特新','新能源','5G','芯片','半导体']
        jjname=pd.DataFrame(data=util.get_mutual_stock_funds('20211231'),columns=['jjdm'])

        sql="select jjdm,jjmc from st_fund.t_st_gm_jjxx "
        jjname=pd.merge(jjname,hbdb.db2df(sql,db='funduser'),how='left',on='jjdm')

        orconditon=[False]*len(jjname)
        for word in key_words:
            orconditon=(orconditon)|(jjname['jjmc'].str.contains(word))

        return  jjname.loc[orconditon]
    @staticmethod
    def advance_man_fromticker(jjdm_list):
        from hbshare.fe.mutual_analysis import holdind_based as hb

        hld=hb.read_hld_fromdb(jjdm_list,start_date='20211231',end_date='20211231')

        ticker_list=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",sheet_name='成份汇总')
        ticker_list['Wind代码']=ticker_list['Wind代码'].astype(str).str.replace("'", "")
        hld=pd.merge(hld,ticker_list,how='left',left_on='zqdm',right_on='Wind代码')
        hld=hld[hld['Wind代码'].notnull()]
        hld=hld.groupby('jjdm').sum()
        hld['zjbl']=hld['zjbl'].rank(method='min')/len(hld)
        hld.rename(columns={'zjbl':'zjbl_ticker'})
        return hld.reset_index()
    @staticmethod
    def advance_ols(jjdm_list):
        from hbshare.fe.mutual_analysis import  nav_based as nb
        import statsmodels.api as sm

        index_range=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\先进制造指数的区间.xlsx")
        index_range['start date'] = index_range['start date'].astype(str).str.replace("-", "")
        index_range['end date'] = index_range['end date'].astype(str).str.replace("-", "")

        index_ret=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",
                                  sheet_name='走势').drop('881001.WI',axis=1)
        index_ret[['930820.CSI','931167.CSI','866003.WI']]=\
            index_ret[['930820.CSI','931167.CSI','866003.WI']].pct_change()
        index_ret['Date']=index_ret['Date'].astype(str).str.replace("-","")
        index_ret=index_ret.iloc[2:]
        index_ret.set_index('Date',inplace=True)
        index_ret = index_ret.sort_index()
        index_ret=index_ret*100

        result=[[],[],[],[],[],[],[],[],[]]
        result2=[[],[],[],[],[],[],[],[],[]]
        result3 = [[], [], [], [], [], [], [], [], []]
        para_df=pd.DataFrame(data=jjdm_list,columns=['jjdm'])

        for jjdm in jjdm_list:
            print(jjdm)
            jj_ret = nb.get_jj_daily_ret([jjdm])
            jj_ret.index=jj_ret.index.astype(str)
            olsdf = pd.merge(index_ret,jj_ret, how='left', left_index=True, right_index=True).fillna(0)

            # laster_3_year = (datetime.datetime.strptime(str(olsdf.index[-1]),
            #                                                                               '%Y%m%d') - datetime.timedelta(
            #     days=365*3)).strftime('%Y%m%d')

            for j in range(3):
                ind=['930820.CSI','931167.CSI','866003.WI'][j]

                #for 正超额
                time_range=index_range[(index_range['index name']==ind)&(index_range['type']=='正超额')]
                tempolsdf=[]
                for i in range(len(time_range)):
                    tempolsdf.append(olsdf.loc[time_range['start date'].iloc[i]:time_range['end date'].iloc[i]])
                tempolsdf=pd.concat(tempolsdf,axis=0)
                #keep only daily ret part,remove weekly return part
                tempolsdf['index_num'] = tempolsdf.reset_index().index.to_list()
                tempolsdf['con']=1
                new_start_date=tempolsdf[tempolsdf[jjdm]==0][tempolsdf[tempolsdf[jjdm]==0]['index_num'].diff()==1]
                if(len(new_start_date)>0):
                    new_start_date=new_start_date.index[-1]
                    tempolsdf=tempolsdf.loc[new_start_date:]
                ols_result=sm.OLS(tempolsdf[ind].values,
                              tempolsdf[[jjdm,'con']].values).fit()
                para =ols_result.params.tolist()
                rsq_adj=ols_result.rsquared_adj
                result[j*3].append(para[0])
                result[(j * 3)+1].append(para[1])
                result[(j * 3)+2].append(rsq_adj)

                #for 负超额
                time_range=index_range[(index_range['index name']==ind)&(index_range['type']=='负超额')]
                tempolsdf=[]
                for i in range(len(time_range)):
                    tempolsdf.append(olsdf.loc[time_range['start date'].iloc[i]:time_range['end date'].iloc[i]])
                tempolsdf=pd.concat(tempolsdf,axis=0)
                #keep only daily ret part,remove weekly return part
                tempolsdf['index_num'] = tempolsdf.reset_index().index.to_list()
                tempolsdf['con']=1
                new_start_date=tempolsdf[tempolsdf[jjdm]==0][tempolsdf[tempolsdf[jjdm]==0]['index_num'].diff()==1]
                if(len(new_start_date)>0):
                    new_start_date=new_start_date.index[-1]
                    tempolsdf=tempolsdf.loc[new_start_date:]
                ols_result=sm.OLS(tempolsdf[ind].values,
                              tempolsdf[[jjdm,'con']].values).fit()
                para =ols_result.params.tolist()
                rsq_adj=ols_result.rsquared_adj
                result3[j*3].append(para[0])
                result3[(j * 3)+1].append(para[1])
                result3[(j * 3)+2].append(rsq_adj)

                # for 负超额+正超额
                time_range = index_range[(index_range['index name'] == ind) ]
                tempolsdf = []
                time_range=time_range.sort_values('start date')
                for i in range(len(time_range)):
                    tempolsdf.append(olsdf.loc[time_range['start date'].iloc[i]:time_range['end date'].iloc[i]])
                tempolsdf = pd.concat(tempolsdf, axis=0)
                tempolsdf.drop_duplicates(ignore_index=False, inplace=True)
                # keep only daily ret part,remove weekly return part
                tempolsdf['index_num'] = tempolsdf.reset_index().index.to_list()
                tempolsdf['con'] = 1
                new_start_date=tempolsdf[tempolsdf[jjdm]==0][tempolsdf[tempolsdf[jjdm]==0]['index_num'].diff()==1]
                if(len(new_start_date)>0):
                    new_start_date=new_start_date.index[-1]
                    tempolsdf=tempolsdf.loc[new_start_date:]

                ols_result=sm.OLS(tempolsdf[ind].values,
                              tempolsdf[[jjdm,'con']].values).fit()
                para =ols_result.params.tolist()
                rsq_adj=ols_result.rsquared_adj
                result2[j*3].append(para[0])
                result2[(j * 3)+1].append(para[1])
                result2[(j * 3)+2].append(rsq_adj)

        for j in range(3):
            ind = ['930820.CSI', '931167.CSI', '866003.WI'][j]

            para_df[ind+'_positive_beta']=result[(j*3)+0]
            para_df[ind + '_positive_alpha'] = result[(j * 3) + 1]
            para_df[ind + '_positive_rsquared_adj'] = result[(j * 3) + 2]

            para_df[ind+'_negative_beta']=result3[(j*3)+0]
            para_df[ind + '_negative_alpha'] = result3[(j * 3) + 1]
            para_df[ind + '_negative_rsquared_adj'] = result3[(j * 3) + 2]

            para_df[ind+'_positive&negative_beta']=result2[(j*3)+0]
            para_df[ind + '_positive&negative_alpha'] = result2[(j * 3) + 1]
            para_df[ind + '_positive&negative_rsquared_adj'] = result2[(j * 3) + 2]


        return  para_df
    @staticmethod
    def advance_contribution(asofdate):

        ticker_list=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",sheet_name='成份汇总')
        ticker_list['Wind代码']=ticker_list['Wind代码'].astype(str).str.replace("'", "")

        #read ticker contribution from local
        sql="SELECT * from hbs_ticker_contribution where asofdate='{}'".format(asofdate)
        ticker_contribution=pd.read_sql(sql,con=localdb)
        #only take top 20 remove last 5
        ticker_contribution['rank']=ticker_contribution.groupby('jjdm').rank(ascending=False)
        ticker_contribution=ticker_contribution[ticker_contribution['rank']<=20].drop('rank',axis=1)


        ticker_contribution=pd.merge(ticker_contribution,ticker_list,how='left',
                                     left_on='zqdm',right_on='Wind代码')
        total_contribution = ticker_contribution.groupby('jjdm').sum()['contribution'].to_frame('total_con')
        ticker_contribution=ticker_contribution[ticker_contribution['Wind代码'].notnull()]
        ticker_contribution=pd.merge(ticker_contribution,total_contribution,how='left',on='jjdm')
        ticker_contribution['contribution_ratio']=ticker_contribution['contribution']/ticker_contribution['total_con']
        ticker_contribution=ticker_contribution.groupby('jjdm').sum()['contribution_ratio'].to_frame('ticker_contribution_ratio')



        advance_man_3ind_cluster=pd.read_excel(r"E:\GitFolder\docs\先进制造三级行业名单20220823.xlsx")[['板块','申万三级']]

        #read 3rd industry contribution from local
        sql="SELECT jjdm,industry_name,contribution from hbs_industry_contribution where industry_lv='sjxymc' and   asofdate='{}' ".format(asofdate)
        industry_contribution=pd.read_sql(sql,con=localdb)
        # only take top 20 remove last 5
        # industry_contribution['rank']=industry_contribution.groupby('jjdm').rank(ascending=False)
        # industry_contribution=industry_contribution[industry_contribution['rank']<=20].drop('rank',axis=1)
        industry_total_con=industry_contribution.groupby('jjdm').sum().rename(columns={'contribution':'total_con'})
        industry_contribution=pd.merge(industry_contribution,industry_total_con,how='left',on='jjdm')
        industry_contribution['contribution_ratio']=industry_contribution['contribution']/industry_contribution['total_con']

        output_df=ticker_contribution.copy()

        for sector in advance_man_3ind_cluster['板块'].unique():

            target_ind = advance_man_3ind_cluster[advance_man_3ind_cluster['板块']==sector]['申万三级'].tolist()
            target_ind=pd.DataFrame(data=target_ind,columns=['industry_name'])
            target_ind['flag']=1
            industry_contribution_temp=pd.merge(industry_contribution,target_ind,how='left',on='industry_name')
            industry_contribution_temp=industry_contribution_temp[industry_contribution_temp['flag'].notnull()]

            industry_contribution_temp=industry_contribution_temp.groupby('jjdm').sum()['contribution_ratio'].to_frame('{}_contribution_ratio'.format(sector))
            output_df=pd.merge(output_df,industry_contribution_temp,how='left',on='jjdm')


        tempdf=industry_contribution.pivot_table('contribution_ratio','jjdm','industry_name').fillna(0)
        tempdf=pd.merge(tempdf
                        ,industry_contribution.pivot_table('contribution','jjdm','industry_name').fillna(0).sum(axis=1).to_frame('total_con')
                        ,how='left',on='jjdm')
        tempdf.columns=[x+'_con' for x in tempdf.columns.tolist()]
        output_df=pd.merge(output_df,tempdf
                           ,how='left',on='jjdm')

        return  output_df

    def get_advance_man(self,asofdate='20211231',filter_pool=False):

        jjdm_list = util.get_mutual_stock_funds(asofdate)

        adv_man_contribution=self.advance_contribution(asofdate)
        adv_man_3ind = self.advance_man_from3ind(asofdate)
        # adv_man_ols=self.advance_ols(jjdm_list)
        # adv_man_ticker=self.advance_man_fromticker(jjdm_list)
        adv_man_words=self.advance_man_wording()

        adv_man=pd.merge(adv_man_3ind,adv_man_words
                         ,how='left',on='jjdm')
        # adv_man=pd.merge(adv_man,adv_man_ticker
        #                  ,how='left',on='jjdm')
        # adv_man=pd.merge(adv_man,adv_man_ols
        #                  ,how='left',on='jjdm')
        adv_man=pd.merge(adv_man,adv_man_contribution
                         ,how='left',on='jjdm')

        if(filter_pool):
            pool=adv_man.copy()
            bmk=util.get_bmk_funds_list('20211231',2)
            pool=pool[pool['jjdm'].isin(bmk)]
            pool=pool.sort_values('先进制造三级行业', ascending=False)[0:200]
            pool=pool.sort_values('total_con_con', ascending=False)[0:100]
            ind3_ret_rank=pd.read_excel(r"E:\GitFolder\docs\三级行业2019至2021年收益排名.xlsx")
            ind3_ret_rank['证券简称']=[x.replace('(申万)','') for x in ind3_ret_rank['证券简称']]
            advance_man_3ind_cluster = pd.read_excel(r"E:\GitFolder\docs\先进制造三级行业名单20220823.xlsx")[['板块','申万三级']]
            ind3_ret_rank=pd.merge(ind3_ret_rank,advance_man_3ind_cluster
                                   ,how='inner',left_on='证券简称',right_on='申万三级')
            top20_ind3=[x+'_w' for x in ind3_ret_rank.sort_values('区间涨跌幅',ascending=False)[0:20]['证券简称']]
            pool['top20_w']=pool[top20_ind3].sum(axis=1)
            pool['top20_w']=pool['top20_w'].rank(method='min')

            for col in advance_man_3ind_cluster['申万三级'].tolist():
                pool[col+'_unit_con']=pool[col+'_con']/pool[col+'_w']

            pool['avg_unit_con']=pool[[x+'_unit_con' for x in advance_man_3ind_cluster['申万三级'].tolist()]].mean(axis=1)
            pool['avg_unit_con'] = pool['avg_unit_con'].rank(method='min')
            pool['new_rank']=pool['avg_unit_con']*0.25+pool['top20_w']*0.75
            # pool=pool.sort_values('new_rank',ascending=False)[0:20][['jjdm']]
            pool['asofdate']=asofdate


            localdb.execute("delete from pool_potient_advance_man where asofdate='{}'".format(asofdate))
            pool[['jjdm','new_rank','asofdate']].to_sql('pool_potient_advance_man',con=localdb,index=False,if_exists='append')

        else:
            adv_man.to_excel('先进制造_基于三级行业列表.xlsx')

            print('Done')

class Similarity_filter:

    @staticmethod
    def theme_similarity(target_jjdm,ext1,ext2):

        maxasofdate=pd.read_sql("select max(asofdate) as max from jjpic{0}_theme_p".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_theme_p where asofdate='{1}' and jjdm='{2}' "\
            .format(ext1,maxasofdate,target_jjdm)
        target_pic=pd.read_sql(sql,con=localdb).fillna(0)

        maxasofdate=pd.read_sql("select max(asofdate) as max from jjpic{0}_theme_p".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_theme_p where asofdate='{1}' "\
            .format(ext2,maxasofdate)
        theme_pic=pd.read_sql(sql,con=localdb).fillna(0)

        theme_weight_difference=((theme_pic['大金融']-target_pic['大金融'].values[0]).abs()\
                                +(theme_pic['消费']-target_pic['消费'].values[0]).abs()+\
                                (theme_pic['TMT']-target_pic['TMT'].values[0]).abs()\
                                +(theme_pic['周期']-target_pic['周期'].values[0]).abs()\
                                +(theme_pic['制造']-target_pic['制造'].values[0]).abs())/2

        theme_difference=((0.1*((theme_pic['主题集中度']-target_pic['主题集中度'].values[0]).abs().rank(method='min')/len(theme_pic))\
                         +0.1*((theme_pic['主题换手率']-target_pic['主题换手率'].values[0]).abs().rank(method='min')/len(theme_pic))\
                         +(theme_weight_difference.rank(method='min')/len(theme_pic))*0.8)/1).to_frame('theme_similarity')
        theme_difference['jjdm']=theme_pic['jjdm']
        theme_difference['theme_similarity']=theme_difference['theme_similarity'].rank(method='min',ascending=False)/len(theme_difference)
        theme_difference['theme_weight_difference']=theme_weight_difference.values

        return  theme_difference

    @staticmethod
    def industry_similarity(target_jjdm,ext1,ext2):

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_p".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_p where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_industry_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_p".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_p where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        industry_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = \
        pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_detail_1".format(ext1), con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_detail_1 where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_industry_detail_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = \
        pd.read_sql("select max(asofdate) as max from jjpic{0}_industry_detail_1".format(ext2), con=localdb)['max'][0]
        sql = "select * from jjpic{0}_industry_detail_1 where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        industry_detail_pic = pd.read_sql(sql, con=localdb).fillna(0)

        industry_weight_difference=pd.merge(industry_detail_pic,
                                            target_industry_detail_pic,
                                            how='outer',on='行业名称').fillna(0)
        industry_weight_difference['weight_difference']=(industry_weight_difference['占持仓比例(时序均值)_x']-industry_weight_difference['占持仓比例(时序均值)_y']).abs()
        industry_weight_difference=industry_weight_difference.groupby('jjdm_x').sum()['weight_difference']/2

        industry_difference=pd.Series(((0.1*((industry_pic['一级行业集中度']-target_industry_pic['一级行业集中度'].values[0]).abs().rank(method='min')/len(industry_pic)).values
                             +0.1*((industry_pic['一级行业换手率']-target_industry_pic['一级行业换手率'].values[0]).abs().rank(method='min')/len(industry_pic)).values
                              +0.8*(industry_weight_difference.rank(method='min') / len(industry_weight_difference)).values)/1)).to_frame('industry_similarity')
        industry_difference['jjdm']=industry_pic['jjdm']
        industry_difference['industry_similarity'] = industry_difference['industry_similarity'].rank(method='min', ascending=False)/len(industry_difference)
        industry_difference['industry_weight_difference']=industry_weight_difference.values
        return  industry_difference

    @staticmethod
    def style_similarity(target_jjdm,ext1,ext2):


        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_value_p_hbs".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_value_p_hbs where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_style_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_value_p_hbs".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_value_p_hbs where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        style_pic = pd.read_sql(sql, con=localdb).fillna(0)

        style_weight_difference = ((style_pic['成长绝对暴露(持仓)'] - target_style_pic['成长绝对暴露(持仓)'].values[0]).abs() \
                                   + (style_pic['价值绝对暴露(持仓)'] - target_style_pic['价值绝对暴露(持仓)'].values[0]).abs() )/2
        style_difference = ((0.1*((style_pic['集中度(持仓)'] - target_style_pic['集中度(持仓)'].values[0]).abs().rank(method='min')/len(style_pic)) \
                            + 0.1*((style_pic['换手率(持仓)'] - target_style_pic['换手率(持仓)'].values[0]).abs().rank(method='min')/len(style_pic)) \
                            + 0.8*(style_weight_difference.rank(method='min')/len(style_pic))) / 1).to_frame('style_similarity')

        style_difference['jjdm']=style_pic['jjdm']
        style_difference['style_similarity']=style_difference['style_similarity'].rank(method='min',ascending=False)/len(style_difference)
        style_difference['style_weight_difference']=style_weight_difference.values

        return style_difference

    @staticmethod
    def size_similarity(target_jjdm,ext1,ext2):

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_size_p_hbs".format(ext1),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_size_p_hbs where asofdate='{1}' and jjdm='{2}' " \
            .format(ext1, maxasofdate, target_jjdm)
        target_size_pic = pd.read_sql(sql, con=localdb).fillna(0)

        maxasofdate = pd.read_sql("select max(asofdate) as max from jjpic{0}_size_p_hbs".format(ext2),con=localdb)['max'][0]
        sql = "select * from jjpic{0}_size_p_hbs where asofdate='{1}' " \
            .format(ext2, maxasofdate)
        size_pic = pd.read_sql(sql, con=localdb).fillna(0)

        size_weight_difference = ((size_pic['大盘绝对暴露(持仓)'] - target_size_pic['大盘绝对暴露(持仓)'].values[0]).abs()
                                  + (size_pic['中盘绝对暴露(持仓)'] - target_size_pic['中盘绝对暴露(持仓)'].values[0]).abs()+
                                  (size_pic['小盘绝对暴露(持仓)'] - target_size_pic['小盘绝对暴露(持仓)'].values[0]).abs() )/2
        size_difference = ((0.1*((size_pic['集中度(持仓)'] - target_size_pic['集中度(持仓)'].values[0]).abs().rank(method='min')/len(size_pic)) \
                            + 0.1*((size_pic['换手率(持仓)'] - target_size_pic['换手率(持仓)'].values[0]).abs().rank(method='min')/len(size_pic)) \
                            + 0.8*(size_weight_difference.rank(method='min')/len(size_pic))) / 1).to_frame('size_similarity')
        size_difference['jjdm']=size_pic['jjdm']
        size_difference['size_similarity']=size_difference['size_similarity'].rank(method='min',ascending=False)/len(size_difference)
        size_difference['size_weight_difference']=size_weight_difference.values

        return size_difference

    def get_similarity(self,target_jjdm,mu_target_type=True,mu_pool_type=True,shown_num=30):

        if(mu_target_type==True):
            ext1=''
        else:
            ext1='_prv'

        if(mu_pool_type==True):
            ext2=''
        else:
            ext2='_prv'


        theme_difference=self.theme_similarity(target_jjdm,ext1,ext2)
        industry_difference=self.industry_similarity(target_jjdm,ext1,ext2)
        style_differene=self.style_similarity(target_jjdm,ext1,ext2)
        size_difference=self.size_similarity(target_jjdm,ext1,ext2)

        similarity=pd.merge(pd.merge(theme_difference,industry_difference,how='inner',on='jjdm'),
                            pd.merge(style_differene,size_difference,how='inner',on='jjdm'),how='inner',on='jjdm')

        similarity['total_similarity']=(1.5*similarity['theme_similarity']+
                                        0.5*similarity['industry_similarity']+
                                        0.5*similarity['style_similarity']+
                                        0.5*similarity['size_similarity'])/3
        similarity=similarity.sort_values('total_similarity',ascending=False).iloc[0:shown_num]

        jj_base_info=hbdb.db2df("select jjdm,jjjc from st_fund.t_st_gm_jjxx where jjdm in ({})"
                                .format(util.list_sql_condition(similarity['jjdm'].tolist())),db='funduser')
        similarity=pd.merge(similarity,jj_base_info,how='left',on='jjdm')

        return similarity[['jjdm','jjjc','total_similarity','theme_similarity','theme_weight_difference',
                           'industry_similarity','industry_weight_difference','style_similarity','style_weight_difference',
                           'size_similarity','size_weight_difference']]

if __name__ == '__main__':

    #sf=Similarity_filter()
    #similarity=sf.get_similarity('006624')
    #similarity = sf.get_similarity('S66391',mu_pool_type=True,mu_target_type=False)
    #similarity.to_excel('similarity_of_{}.xlsx'.format('006624'),index=False)

    pe=Positive_equilibrium()
    pe_pool=[]
    for asofdate in ['20151231', '20160630', '20161231', '20170630', '20171231', '20180630'
        , '20181231', '20190630', '20191231'
        , '20200630', '20201231', '20210630', '20211231', '20220630']:

        tempdf=pd.DataFrame()
        tempdf['jjdm']=pe.get_nav_equilibrium_pool(asofdate)
        tempdf['asofdate']=asofdate
        pe_pool.append(tempdf)
    pd.concat(pe_pool,axis=0).to_excel('被动均衡（净值）池.xlsx')



    sql="select ticker,category,vcg_score,trade_date,type from st_fund.r_st_gp_equity_score where trade_date>='20200630'    "
    test=hbdb.db2df(sql,db='funduser')
    test=test.sort_values(['ticker','trade_date','type'])
    test=test.drop_duplicates(['ticker','trade_date'],keep='first')
    test['size_type']=test['category'].str[0:2]
    test['style_type']='成长'
    test.loc[test['vcg_score']<=150,'style_type'] = '价值'
    test['category']=test['size_type']+test['style_type']
    reprot_date=test['trade_date'].unique().tolist()


    outdf=[]

    for i in range(len(reprot_date)):

        if(i!=len(reprot_date)-1):
            t0 = reprot_date[i]
            t1 = reprot_date[i + 1]
            sql = """ select TDATE,SYMBOL,VATURNOVER from finchina.CHDQUOTE 
            where TDATE>{0} and TDATE<={1} and SYMBOL='000001'
            """.format(t0,t1)
            date_list=hbdb.db2df(sql, db='readonly').drop('ROW_ID', axis=1)['TDATE'].tolist()
        else:
            t0 = reprot_date[i]
            sql = """ select TDATE,SYMBOL,VATURNOVER from finchina.CHDQUOTE 
            where TDATE>{0}  and SYMBOL='000001'
            """.format(t0)
            date_list=hbdb.db2df(sql, db='readonly').drop('ROW_ID', axis=1)['TDATE'].tolist()

        for date in date_list:
            print(date)
            sql = """ select TDATE,SYMBOL,VATURNOVER from finchina.CHDQUOTE 
            where TDATE={0} 
            """.format(date)
            cjl = hbdb.db2df(sql, db='readonly').drop('ROW_ID', axis=1)
            cjl=pd.merge(cjl,test[test['trade_date']==t0]
                         ,how='left',left_on='SYMBOL',right_on='ticker')
            outdf.append(cjl.groupby(['category', 'TDATE']).sum()['VATURNOVER'].reset_index())

    outdf=pd.concat(outdf,axis=0)
    outdf=outdf.pivot_table('VATURNOVER','TDATE','category')
    cols=outdf.columns.tolist()
    outdf['total']=outdf.sum(axis=1)
    for col in cols:
        outdf[col]=outdf[col]/outdf['total']

    outdf.to_excel('成交量分布.xlsx')
    plot= functionality.Plot(1200,600)
    plot.plotly_area(outdf.drop('total',axis=1),'成交量分布')

    # ir=Industry_rotating()
    # industry_roation=ir.get_industry_rotating(shift_count_rank_thresheld=0.6
    #                           ,shift_ret_thresheld=0.6,shift_winpro_thresheld=0.6
    #                                           ,left_trade_ratio=0.4,pool_size=20)

    # val=Value()
    # growth_pool,value_pool=val.get_value(show_num=200, shift_ratio_threshold=0.5)
    # growth_pool.to_excel(r"E:\GitFolder\docs\金工基金池\成长池.xlsx",index=False)
    # value_pool.to_excel(r"E:\GitFolder\docs\金工基金池\价值池.xlsx",index=False)


    # am=Advance_man()
    # am.get_advance_man(asofdate='20211231',filter_pool=True)



    print('')
    #


    # rp=Report()
    # rp.filter_pool_report()

    #advance_man()


    #
    # equclass=Equilibrium()
    # ind_equ, style_equ,size_equ, \
    # pepbroe_equ, nav_equ,joint_rank, \
    # joint_result=equclass.get_equilibrium(threshield=0.3,show_num=100)
    #
    #
    #
    # plot=functionality.Plot(800,800)
    #
    # plot.plotly_table(ind_equ, 800, 'asdf')
    # pla.pool_picturing(ind_equ)
    #
    # plot.plotly_table(style_equ, 800, 'asdf')
    # pla.pool_picturing(style_equ)
    #
    # plot.plotly_table(size_equ, 800, 'asdf')
    # pla.pool_picturing(size_equ)
    #
    # plot.plotly_table(pepbroe_equ, 800, 'asdf')
    # pla.pool_picturing(pepbroe_equ)
    #
    # plot.plotly_table(nav_equ, 800, 'asdf')
    # pla.pool_picturing(nav_equ)
    #
    # plot.plotly_table(joint_result,800,'asdf')
    # pla.pool_picturing(joint_result)
    #
    # plot.plotly_table(joint_rank, 800, 'asdf')
    # pla.pool_picturing(joint_rank)

    #

    # leftclass=Leftside()
    #
    # stock_left, ind_left, value_left, \
    # size_left, joint_rank, joint_restult=leftclass.get_left(threshield=0.4,show_num=100)
    #
    #
    # plot=functionality.Plot(800,800)
    #
    # plot.plotly_table(stock_left,800,'asdf')
    # pla.pool_picturing(stock_left)
    #
    # plot.plotly_table(ind_left, 800, 'asdf')
    # pla.pool_picturing(ind_left)
    #
    # plot.plotly_table(value_left, 800, 'asdf')
    # pla.pool_picturing(value_left)
    #
    # plot.plotly_table(size_left, 800, 'asdf')
    # pla.pool_picturing(size_left)
    #
    # plot.plotly_table(joint_rank, 800, 'asdf')
    # pla.pool_picturing(joint_rank)

    # plot.plotly_table(joint_restult, 800, 'asdf')
    # pla.pool_picturing(joint_restult)

    # sizeclass = Size()
    # big_size, medium_size, small_size = sizeclass.get_size(fre='M', show_num=100,
    #                                                        shift_ratio_threshold=0.5,
    #                                                        centralization_threshold=0.5)

    #plot = functionality.Plot(800, 800)
    # plot.plotly_table(big_size, 800, 'asdf')
    # pla.pool_picturing(big_size)
    #
    # plot.plotly_table(medium_size, 800, 'asdf')
    # pla.pool_picturing(medium_size)
    #
    # plot.plotly_table(small_size, 800, 'asdf')
    # pla.pool_picturing(small_size)


    # valueclass = Value()
    # growth, value=valueclass.get_value(fre='M', show_num=20, shift_ratio_threshold=0.5, centralization_threshold=0.5)
    #     , absolute_pe_value, absolute_pb_value, relative_pe_value, relative_pb_value, \
    # dividend_value, reverse_value, high_quality_value \


    # plot = functionality.Plot(800, 800)
    # plot.plotly_table(growth, 800, 'asdf')
    # pla.pool_picturing(growth)
    #
    # plot.plotly_table(value, 800, 'asdf')
    # pla.pool_picturing(value)
    #
    # plot.plotly_table(absolute_pe_value, 800, 'asdf')
    # pla.pool_picturing(absolute_pe_value)
    #
    # plot.plotly_table(absolute_pb_value, 800, 'asdf')
    # pla.pool_picturing(absolute_pb_value)
    #
    # plot.plotly_table(relative_pe_value, 800, 'asdf')
    # pla.pool_picturing(relative_pe_value)
    #
    # plot.plotly_table(relative_pb_value, 800, 'asdf')
    # pla.pool_picturing(relative_pb_value)
    #
    # plot.plotly_table(dividend_value, 800, 'asdf')
    # pla.pool_picturing(dividend_value)
    #
    # plot.plotly_table(reverse_value, 800, 'asdf')
    # pla.pool_picturing(reverse_value)
    #
    # plot.plotly_table(high_quality_value, 800, 'asdf')
    # pla.pool_picturing(high_quality_value)
    #
    # print("Done")


